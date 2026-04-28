#!/usr/bin/env python3
"""
AmarSheba - Automated Viewport Screenshot & Documentation Generator
====================================================================
Generates MD, DOCX, and a single PORTABLE HTML file.

Usage:
  python docs/screenshot_docs.py
"""

import base64
import datetime
import os
import re
from pathlib import Path

from playwright.sync_api import Page, sync_playwright

try:
    from docx import Document
    from docx.shared import Inches
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# --- Configuration ------------------------------------------------------------
BASE_URL = "http://localhost:5174"
VIEWPORTS = [
    {"name": "Desktop", "width": 1440, "height": 900},
    {"name": "Tablet",  "width": 768,  "height": 1024},
    {"name": "Mobile",  "width": 390,  "height": 844},
]

ROLE_PAGES = [
    {"role": "public", "slug": "home", "path": "/", "title": "Home"},
    {"role": "public", "slug": "services", "path": "/services", "title": "Services"},
    {"role": "public", "slug": "find-providers", "path": "/find", "title": "Find Providers"},
    {"role": "public", "slug": "about", "path": "/about", "title": "About"},
    {"role": "public", "slug": "access", "path": "/access", "title": "Role Access"},
    {"role": "customer", "slug": "cust-dashboard", "path": "/dashboard", "title": "Customer Dashboard"},
    {"role": "provider", "slug": "prov-dashboard", "path": "/provider-app", "title": "Provider Dashboard"},
    {"role": "resource", "slug": "res-home", "path": "/resource-app", "title": "Resource Home"},
    {"role": "admin", "slug": "admin-dashboard", "path": "/admin", "title": "Admin Dashboard"},
]

SCREENSHOTS_DIR = Path(__file__).parent / "screenshots"
MD_OUTPUT       = Path(__file__).parent / "PROJECT_DOCS.md"
DOCX_OUTPUT     = Path(__file__).parent / "PROJECT_DOCS.docx"
HTML_OUTPUT     = Path(__file__).parent / "PROJECT_DOCS.html"

WAIT_MS = 2200

ROLE_DESCRIPTIONS = {
    "public":   ("Public / Guest", "No auth required."),
    "customer": ("Customer", "Registered users booking services."),
    "provider": ("Provider", "Verified service providers."),
    "resource": ("Resource", "Field staff workers."),
    "admin":    ("Admin", "Platform administrators."),
}

TECH_STACK = [
    ("Frontend", "React 18, Tailwind CSS v4"),
    ("Backend", "Laravel 11, MySQL 8"),
    ("Automation", "Playwright, Python"),
]

# --- Helpers ------------------------------------------------------------------

def slugify(text: str) -> str:
    return re.sub(r"[^a-z0-9]+", "_", text.lower()).strip("_")

def screenshot_path(viewport_name: str, role: str, page_slug: str) -> Path:
    out_dir = SCREENSHOTS_DIR / slugify(viewport_name) / role
    out_dir.mkdir(parents=True, exist_ok=True)
    return out_dir / f"{page_slug}.png"

def sign_in_as(page: Page, role: str):
    if role == "public": return
    try:
        page.goto(f"{BASE_URL}/access", wait_until="domcontentloaded", timeout=15000)
        label = {"customer": "Customer", "provider": "Provider", "resource": "Resource", "admin": "Admin"}[role]
        page.get_by_role("button", name=re.compile(label, re.IGNORECASE)).first.click()
        page.wait_for_timeout(1000)
    except: pass

def capture_page(page: Page, url: str, out_path: Path):
    try:
        page.goto(url, wait_until="domcontentloaded", timeout=15000)
        page.wait_for_timeout(WAIT_MS)
        page.screenshot(path=str(out_path), full_page=True)
        print(f"    [OK] {out_path.name}")
        return True
    except: return False

# --- Documentation Generators -------------------------------------------------

def build_markdown(results):
    lines = ["# AmarSheba Documentation\n"]
    for slug, vp_map in results.items():
        pg = next(p for p in ROLE_PAGES if p["slug"] == slug)
        lines.append(f"## {pg['title']} ({pg['role']})\nRoute: `{pg['path']}`\n")
        for vp_name, path in vp_map.items():
            if path:
                rel = str(path.relative_to(MD_OUTPUT.parent)).replace("\\", "/")
                lines.append(f"### {vp_name}\n![{vp_name}]({rel})\n")
    MD_OUTPUT.write_text("\n".join(lines), encoding="utf-8")

def build_docx(results):
    if not HAS_DOCX: return
    doc = Document()
    doc.add_heading('AmarSheba Documentation', 0)
    for slug, vp_map in results.items():
        pg = next(p for p in ROLE_PAGES if p["slug"] == slug)
        doc.add_heading(f"{pg['title']} ({pg['role']})", level=1)
        for vp_name, path in vp_map.items():
            if path:
                doc.add_heading(vp_name, level=2)
                doc.add_picture(str(path), width=Inches(6))
    doc.save(str(DOCX_OUTPUT))

def build_html(results):
    now = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
    html = [f"""<!DOCTYPE html><html><head><title>AmarSheba Docs</title>
    <style>
        body {{ font-family: sans-serif; max-width: 1000px; margin: 40px auto; padding: 20px; background: #f4f7f9; }}
        .card {{ background: white; padding: 30px; border-radius: 15px; box-shadow: 0 4px 20px rgba(0,0,0,0.05); margin-bottom: 40px; }}
        h1 {{ color: #004ac6; }}
        h2 {{ border-bottom: 2px solid #eee; padding-bottom: 10px; margin-top: 40px; }}
        img {{ max-width: 100%; border-radius: 8px; border: 1px solid #ddd; margin-top: 10px; }}
        .meta {{ color: #666; font-size: 0.9em; }}
        .vp-grid {{ display: grid; gap: 20px; margin-top: 20px; }}
    </style></head><body>
    <h1>AmarSheba Project Documentation</h1>
    <p class="meta">Generated on {now}</p>"""]
    
    for slug, vp_map in results.items():
        pg = next(p for p in ROLE_PAGES if p["slug"] == slug)
        html.append(f"<div class='card'><h2>{pg['title']} ({pg['role']})</h2><p>Route: <code>{pg['path']}</code></p>")
        for vp_name, path in vp_map.items():
            if path:
                with open(path, "rb") as f:
                    b64 = base64.b64encode(f.read()).decode()
                html.append(f"<h3>{vp_name}</h3><img src='data:image/png;base64,{b64}'>")
        html.append("</div>")
    
    html.append("</body></html>")
    HTML_OUTPUT.write_text("\n".join(html), encoding="utf-8")

# --- Main ---------------------------------------------------------------------

def main():
    SCREENSHOTS_DIR.mkdir(parents=True, exist_ok=True)
    results = {}
    with sync_playwright() as pw:
        browser = pw.chromium.launch(headless=True)
        for vp in VIEWPORTS:
            print(f"\nViewport: {vp['name']}")
            context = browser.new_context(viewport={"width": vp["width"], "height": vp["height"]})
            page = context.new_page()
            current_role = None
            for pg in ROLE_PAGES:
                if pg["role"] != current_role:
                    sign_in_as(page, pg["role"])
                    current_role = pg["role"]
                path = screenshot_path(vp["name"], pg["role"], pg["slug"])
                capture_page(page, BASE_URL + pg["path"], path)
                if pg["slug"] not in results: results[pg["slug"]] = {}
                results[pg["slug"]][vp["name"]] = path
            context.close()
        browser.close()
    
    print("\nBuilding documents...")
    build_markdown(results)
    build_docx(results)
    build_html(results)
    print(f"Done!\nMD: {MD_OUTPUT}\nDOCX: {DOCX_OUTPUT}\nHTML: {HTML_OUTPUT}")

if __name__ == "__main__":
    main()
